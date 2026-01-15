"""One-off generator for anima v2 Deep-Dive.docx — safe to delete after run."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "anima-v2-deep-dive.docx"


def bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def num_list(doc: Document, items: list[str]) -> None:
    for t in items:
        doc.add_paragraph(t, style="List Number")


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("anima v2 — Deep Dive (ELI5 + Algorithms)", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Audience: you. Goal: know exactly what you are building so you can explain it to engineers, "
        "researchers, and interviewers — without claiming a language model has feelings."
    )
    p = doc.add_paragraph()
    run = p.add_run("Hard rule: ")
    run.bold = True
    p.add_run(
        "Say dimensional readout, internal geometry, brain-aligned coordinate, population-level analogy. "
        "Never say “the model feels happy.”"
    )
    doc.add_paragraph("Saved path: " + str(OUT))

    doc.add_page_break()

    doc.add_heading("1. What problem are you solving?", level=1)
    doc.add_paragraph(
        "People treat LLM output text as if it were the whole story. Under the hood, a transformer "
        "maintains a high-dimensional internal state while it chooses the next token. That internal "
        "state can disagree with the eventual words — similar in spirit to cognitive science ideas "
        "about hidden processing versus expressed behavior."
    )
    doc.add_paragraph("Your project builds a pipeline that:")
    num_list(
        doc,
        [
            "Reads hidden activations from several layers during inference.",
            "Optionally relates those signals to human brain recordings (fMRI) using encoding-model ideas.",
            "Trains small linear probes that map activations to three numbers you expose in a UI: "
            "valence (pleasant/unpleasant axis), arousal (activated/calm axis), and uncertainty "
            "(how ambiguous the next-token distribution looks).",
            "Flags large disagreements between early-layer and late-layer probe readouts as a "
            "dashboard alert — framed as internal inconsistency, not “lying emotions.”",
        ],
    )
    doc.add_paragraph(
        "Important humility: this does not prove anything about consciousness. It is instrumentation."
    )

    doc.add_heading("1.1 Mental model: three different objects", level=2)
    bullets(
        doc,
        [
            "Observed text — what you can read.",
            "Hidden activations — tensors inside the network.",
            "Brain BOLD signals — slow, noisy proxy of neural activity averaged over large populations "
            "of neurons and seconds of time.",
        ],
    )
    doc.add_paragraph(
        "Those three are related mathematically only after you make explicit alignment choices. "
        "Most mistakes come from smuggling metaphors across these boundaries."
    )

    doc.add_heading("2. Transformer in one coffee cup", level=1)
    doc.add_paragraph(
        "Input text becomes token IDs. Each token gets an embedding vector. A stack of layers updates "
        "a residual stream: roughly, at each layer you add a learned update to the running representation."
    )
    doc.add_paragraph(
        "At the end, a linear map turns the final hidden vector into logits — one score per vocabulary "
        "token. Softmax turns logits into probabilities for the next token."
    )

    doc.add_heading("2.1 What you hook", level=2)
    doc.add_paragraph(
        "Your spec hooks selected transformer blocks (early-mid, mid, late-mid, late — exact indices "
        "per model family live in layer_config). After a layer runs, you grab output[0]: shape "
        "[batch, sequence_length, hidden_dim]. For generation you usually care about the last position "
        "row — the representation right before predicting the next token."
    )

    doc.add_heading("3. Forward hooks (Phase 1 core)", level=1)
    doc.add_paragraph(
        "PyTorch lets you attach a function that runs after a module’s forward pass. You register it "
        "on each chosen layer. Each forward clears a buffer, runs the model, fills the buffer with "
        "detached CPU tensors (to avoid GPU memory blowups), then you read last-token slices."
    )

    doc.add_heading("3.1 Lifecycle traps", level=2)
    bullets(
        doc,
        [
            "Always remove handles after a session — leaked hooks keep tensors alive.",
            "If you remove hooks at the end of extract(), you must re-register before the next extract(); "
            "otherwise later calls silently fail or read stale buffers.",
            "Different HF architectures expose layers differently (Llama uses model.model.layers; "
            "others differ). Your config must match reality.",
        ],
    )

    doc.add_heading("3.2 Token-wise generation loop (why it is slow)", level=2)
    doc.add_paragraph(
        "Your reference extractor disables KV cache (use_cache=False) and runs output_attentions=True. "
        "That recomputes the whole sequence every step — correct for simple correctness, expensive at "
        "scale. Production systems usually cache keys/values and optionally sample attentions sparsely."
    )

    doc.add_heading("4. Uncertainty signals (three ingredients → one scalar)", level=1)

    doc.add_heading("4.1 Softmax entropy", level=2)
    doc.add_paragraph(
        "Given logits vector z of length V (vocab size), probabilities p = softmax(z). "
        "Entropy H = −Σ_i p_i log(p_i). Normalize by log(V) so rough range is [0,1]. "
        "High entropy ≈ flat distribution ≈ many plausible next tokens."
    )

    doc.add_heading("4.2 Logit gap (top-1 minus top-2)", level=2)
    doc.add_paragraph(
        "Sort logits; take gap g = z_(1) − z_(2). Large gap means one winner dominates. "
        "Your spec maps gap to an uncertainty-like scalar using u_gap = 1 / (1 + max(g, 0)). "
        "Interpretation: confident sampling → smaller uncertainty from this term."
    )

    doc.add_heading("4.3 Attention entropy (last layer, last query position)", level=2)
    doc.add_paragraph(
        "If attentions returned, take the last decoder layer’s attention weights for the final token "
        "attending to all prior positions. For each head compute entropy over the attended positions; "
        "average across heads; normalize by log(sequence_length). Diffuse attention bumps this term."
    )

    doc.add_heading("4.4 Fusion", level=2)
    doc.add_paragraph(
        "Your prototype uses a fixed convex combination — weights sum to 1 so the fused score stays "
        "interpretable as a blended belief:"
    )
    doc.add_paragraph(
        "u_fused = 0.35 · H_norm + 0.35 · u_gap + 0.30 · H_attn",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "This fused value can supervise an uncertainty head on the probe (cheap signal without fMRI)."
    )

    doc.add_heading("5. Brain side — Narratives fMRI in plain English", level=1)
    doc.add_paragraph(
        "Functional MRI measures slow changes in blood oxygen (BOLD). When brain regions work harder, "
        "blood flow changes — lagged and smeared out over seconds. One snapshot every TR (here ~2 s). "
        "Each voxel traces a noisy time series."
    )
    doc.add_paragraph(
        "The Narratives dataset pairs stories with many subjects listening inside the scanner. "
        "You want to predict voxel activity from features derived from the words heard — classic "
        "encoding modeling."
    )

    doc.add_heading("5.1 Hemodynamic lag (HRF)", level=2)
    doc.add_paragraph(
        "Neural response peaks seconds after the stimulus. A blunt first fix is shifting fMRI forward "
        "by a few TRs relative to stimulus-aligned features so peaks line up better. Better models convolve "
        "with an HRF kernel or learn finite impulse response weights — your spec uses trimming/shift; "
        "document that as simplified."
    )

    doc.add_heading("6. Aligning LLM vectors to TR windows", level=1)
    doc.add_paragraph(
        "Words arrive with onsets in seconds. Divide time by TR to bucket each word into a TR index. "
        "Within each TR bin, average hidden vectors for all tokens whose words fall there — your code "
        "template aligns activations list index with word timing index; that only works if tokenization "
        "lines up with words (subword models break naive alignment). You must define aggregation over "
        "subword pieces back to words before binning."
    )

    doc.add_heading("7. Confounds — why word rate wrecks naive claims", level=1)
    doc.add_paragraph(
        "Brains track stimulus dynamics: bursts of words versus silence change BOLD even if semantics "
        "stay boring. Encoding studies often partial out word rate and slow drift (linear + quadratic "
        "time). Your residualize step fits Ridge: confounds → features, subtract prediction."
    )

    doc.add_heading("8. Ridge regression encoding model", level=1)
    doc.add_paragraph(
        "Let X ∈ ℝ^{T×D} be aligned LLM features per TR (possibly residualized). "
        "Let Y ∈ ℝ^{T×V} be fMRI voxels. Ridge minimizes:"
    )
    doc.add_paragraph(
        "||Y − XW||²_F + α ||W||²_F",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Use sklearn.linear_model.RidgeCV with temporal splits — beware autocorrelation when interpreting "
        "standard CV. Evaluate per-voxel Pearson r between predicted and held-out actual time series."
    )
    doc.add_paragraph(
        "High correlation means “this linear map from LLM features explains variance,” not “the brain "
        "computes the LLM.”"
    )

    doc.add_heading("9. Turning brain maps into probe targets (danger zone)", level=1)
    doc.add_paragraph(
        "Your scaffold averages voxel ranges standing in for amygdala / vmPFC / ACC. Without atlas masks "
        "in the same space as preprocessed BOLD, those indices are meaningless — interviewers will "
        "challenge this if presented as neuroscience."
    )
    bullets(
        doc,
        [
            "Fix A (honest): PCA compressed cortical targets — no anatomy claims until masks exist.",
            "Fix B (better): Harvard-Oxford / Schaefer parcellation masks registered to your preprocessed space.",
        ],
    )
    doc.add_paragraph(
        "Literature shortcuts like “amygdala − vmPFC drives valence” are contested and context-dependent; "
        "treat any scalar summary as a pragmatic coordinate, not ground truth."
    )

    doc.add_heading("10. Linear probes (Phase 3)", level=1)
    doc.add_paragraph(
        "A probe is a tiny classifier/regressor on top of frozen activations. Here you learn softmax "
        "weights over layers, average layer vectors into one fused hidden vector h, then apply three "
        "linear heads:"
    )
    bullets(
        doc,
        [
            "valence ≈ tanh(w_v · h + b_v) ∈ (−1,1)",
            "arousal ≈ σ(w_a · h + b_a) ∈ (0,1)",
            "uncertainty_head ≈ σ(w_u · h + b_u) ∈ (0,1), optionally Platt-calibrated against fused logits uncertainty",
        ],
    )
    doc.add_paragraph(
        "Training mixes losses (MSE typical). Temporal splits only — never shuffle time."
    )

    doc.add_heading("10.1 Platt scaling intuition", level=2)
    doc.add_paragraph(
        "Sigmoid outputs might be miscalibrated as probabilities. Platt fits scalar a,b so σ(a·logit + b) "
        "matches targets — with continuous uncertainty you can stick to MSE or bin edges carefully."
    )

    doc.add_heading("11. Suppression / layer disagreement", level=1)
    doc.add_paragraph(
        "Apply the probe to early-layer vs late-layer activations. Large positive Δ valence from "
        "early→late flags valence_suppression in the prototype; large negative Δ uncertainty flags "
        "uncertainty_overclaim."
    )
    doc.add_paragraph(
        "Interview-safe wording: “We measure inconsistent affect geometry across depth.” Avoid claiming "
        "vendor-specific replication unless protocols match."
    )

    doc.add_heading("12. API & dashboard contracts", level=1)
    bullets(
        doc,
        [
            "REST batch vs WebSocket streaming — ensure streamed payloads carry suppression flags if the UI expects them.",
            "Region analog strings must begin with the mandated psychology disclaimer sentence.",
            "Uncertainty above threshold routes to “no human analog” label — avoids fake limbic storytelling.",
        ],
    )

    doc.add_heading("13. Evaluation discipline (what skeptics ask)", level=1)
    bullets(
        doc,
        [
            "Word-rate-only baseline encoding.",
            "Shallow lexical baseline (bag-of-words, TF-IDF).",
            "Temporal cross-validation respecting autocorrelation.",
            "Probe drift across checkpoints / quantization.",
            "Steering sanity checks with random directions as controls.",
        ],
    )

    doc.add_heading("14. Interview cheat sheet", level=1)

    doc.add_heading("14.1 Thirty-second pitch", level=2)
    doc.add_paragraph(
        "“We instrument transformer internals during decoding — multi-layer residual vectors — and fuse "
        "classic distributional uncertainty cues. Separately, we can regularize probes using human fMRI "
        "encoding objectives on public listening data, with strict confound controls and temporal splits. "
        "The UI exposes low-dimensional coordinates with careful language: readouts and analogies, not feelings.”"
    )

    doc.add_heading("14.2 Likely follow-ups + tight answers", level=2)
    doc.add_paragraph("Q: Does high correlation mean the brain runs GPT?", style="Heading 3")
    doc.add_paragraph(
        "A: No — it means linearly predictable shared variance under this stimulus setup and preprocessing."
    )
    doc.add_paragraph("Q: Why ridge not deep decoder?", style="Heading 3")
    doc.add_paragraph(
        "A: Sample efficiency, stability, standard baseline in encoding literature — deep comes later with care."
    )
    doc.add_paragraph("Q: Biggest validity threat?", style="Heading 3")
    doc.add_paragraph(
        "A: Misaligned tokens-to-words-to-TR and bogus ROI masks — fix before scientific claims."
    )
    doc.add_paragraph("Q: What is suppression measuring?", style="Heading 3")
    doc.add_paragraph(
        "A: Disagreement between shallow vs deep probe projections — heuristic stability metric, not moral judgment."
    )

    doc.add_paragraph()
    p = doc.add_paragraph("End of document.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    main()
